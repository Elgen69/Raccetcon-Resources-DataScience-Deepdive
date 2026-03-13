"""
Generate Team CFM - FreshCatch VPD DOCX
HackUSC Hackathon — Semester Long
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# --- Page Margins ---
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

# --- Helper Functions ---
def styled_heading(doc, text, level=1, center=True, size=12, bold=True, italic=False, color=None):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def body_para(doc, text, size=11, justify=True, bold=False, italic=False, space_before=0, space_after=4):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def section_title(doc, number, title, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{number}.     {title.upper()}")
    run.bold = False
    run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p

def subsection_title(doc, letter, title, size=11):
    p = doc.add_paragraph()
    run = p.add_run(f"{letter}  {title}")
    run.bold = False
    run.italic = False
    run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    return p

# =============================================
# TITLE
# =============================================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("FreshCatch")
r.font.size = Pt(18)
r.bold = False

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("by Team CFM — Cebu Fish Market")
r2.font.size = Pt(12)
r2.bold = True

doc.add_paragraph()

# =============================================
# MEMBER INFO — Table of 3 columns
# =============================================
members = [
    ("Elgen Subar", "Department of Computer Science", "BS Computer Science — 3rd Year", "elgen@email.com"),
    ("[Member 2 Name]", "Department of Computer Science", "BS Computer Science — [Year]", "member2@email.com"),
    ("[Member 3 Name]", "Department of Computer Science", "BS Computer Science — [Year]", "member3@email.com"),
]

member_table = doc.add_table(rows=1, cols=3)
member_table.style = 'Table Grid'
# Remove borders
for cell in member_table.rows[0].cells:
    for border in ['top', 'left', 'bottom', 'right']:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        bd = OxmlElement(f'w:{border}')
        bd.set(qn('w:val'), 'none')
        tcBorders.append(bd)
        tcPr.append(tcBorders)

for i, (name, dept, year, email) in enumerate(members):
    cell = member_table.rows[0].cells[i]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.font.size = Pt(10)
    for line in [dept, year, email]:
        p2 = cell.add_paragraph(line)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.runs[0].font.size = Pt(9)
        p2.runs[0].italic = True

doc.add_paragraph()

# =============================================
# ABSTRACT
# =============================================
abstract_table = doc.add_table(rows=1, cols=2)
abstract_table.style = 'Table Grid'

# Remove all borders 
for cell in abstract_table.rows[0].cells:
    for border in ['top', 'left', 'bottom', 'right']:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        bd = OxmlElement(f'w:{border}')
        bd.set(qn('w:val'), 'none')
        tcBorders.append(bd)
        tcPr.append(tcBorders)

left_cell = abstract_table.rows[0].cells[0]
right_cell = abstract_table.rows[0].cells[1]

# Left col — Abstract text
lp = left_cell.paragraphs[0]
lp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
lr_bold = lp.add_run("Abstract")
lr_bold.bold = True
lr_bold.italic = True
lr_bold.font.size = Pt(10)
lr = lp.add_run(
    "—FreshCatch is a B2B digital seafood coordination platform that directly connects "
    "Cebu's small-scale coastal fishermen with restaurants, resorts, and hotels. "
    "Built around the concept of digital Bayanihan, FreshCatch eliminates the information "
    "asymmetry that traps fishermen in exploitative middleman supply chains. "
    "By enabling restaurants to post advance seafood demand and fishermen to see committed "
    "buyers before sailing, the platform reduces seafood waste, stabilizes fisherman income, "
    "and creates a transparent, efficient seafood ecosystem for Cebu's coastal communities. "
    "The MVP focuses on demand coordination—no logistics replacement required— "
    "making the solution buildable and deployable within weeks."
)
lr.font.size = Pt(10)
lr.italic = True

# Right col — Beneficiaries intro
rp = right_cell.paragraphs[0]
rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
rr = rp.add_run(
    "FreshCatch's primary beneficiaries are small-scale fishermen in Cebu's coastal "
    "communities—particularly in Cordova, Lapu-Lapu, Bantayan, Medellin, and Daanbantayan. "
    "Secondary users include restaurants, resorts, and hotels sourcing local seafood daily. "
    "Tertiary stakeholders include fishing cooperatives, local government units, and tourism "
    "establishments invested in sustainable sourcing and community livelihood development."
)
rr.font.size = Pt(10)

doc.add_paragraph()

# =============================================
# I. PROBLEM STATEMENT
# =============================================
section_title(doc, "I.", "Problem Statement")
body_para(doc,
    "Cebu is an island province deeply reliant on fishing as both a cultural identity and an economic backbone. "
    "Despite the abundance of maritime resources and the province's status as a tourism and culinary destination, "
    "small-scale fishermen remain economically marginalized. The root cause is not a lack of supply—it is a "
    "systemic failure in information flow and market access."
)
body_para(doc,
    "The current seafood supply chain in Cebu operates through a middleman-dominated 'bagsakan' system: "
    "fishermen sell their catch to middlemen at suppressed prices, who then sell to restaurants at a significant "
    "markup. This forces fishermen to accept below-market rates, restaurants to pay inflated costs, and the "
    "entire ecosystem to bear unnecessary seafood waste because catch supply is never aligned with buyer demand "
    "in advance. Fishermen sail without guaranteed buyers; restaurants face inconsistent supply with no price "
    "transparency. The digital gap between both sides empowers intermediaries and disempowers the producers."
)
body_para(doc,
    "This problem affects entire coastal barangays whose livelihoods depend on equitable seafood market access—"
    "a crisis that is solvable through a targeted digital coordination layer. Market data on local fish pricing, "
    "catch volumes, and buyer demand are non-existent in accessible formats for fishermen, creating a structural "
    "information asymmetry that perpetuates poverty cycles in Cebu's fishing communities."
)

# =============================================
# II. VALUE PROPOSITION
# =============================================
section_title(doc, "II.", "Value Proposition")
body_para(doc,
    "FreshCatch delivers value by solving the core information gap in Cebu's seafood supply chain—without "
    "requiring physical logistics overhaul. The platform's value is three-dimensional: it empowers fishermen "
    "economically, stabilizes restaurant supply chains, and contributes to sustainable local resource use."
)
body_para(doc,
    "For fishermen: FreshCatch gives them a direct window into buyer demand before they set sail. They post "
    "their expected catch—fish type, volume, landing date, and price—and can see committed restaurant orders in "
    "real time. This eliminates blind selling to middlemen, guarantees a buyer, and allows price negotiation "
    "based on actual market rates. The commitment system transforms fishing from a speculative activity to a "
    "demand-driven livelihood."
)
body_para(doc,
    "For restaurants: FreshCatch provides reliable, advance access to local seafood sourcing. Restaurant "
    "procurement teams can post weekly or daily demand, browse fisherman listings, and commit to purchase "
    "volumes days in advance. This reduces supply uncertainty, eliminates dependence on middlemen, and enables "
    "verified locally sourced seafood claims—a growing differentiator in Cebu's tourism-driven food economy."
)
body_para(doc,
    "For the Cebuano community: FreshCatch embodies Bayanihan in Bytes. It digitizes community cooperation—"
    "connecting the sea to the table through transparency, trust, and mutual benefit. Every transaction that "
    "bypasses an exploitative intermediary is a direct reinvestment into a fisherman's family and a coastal "
    "community's future. FreshCatch aligns economic incentives with community empowerment."
)

# =============================================
# III. TARGET BENEFICIARIES
# =============================================
section_title(doc, "III.", "Target Beneficiaries")

body_para(doc, "Primary Users — Small-Scale Fishermen", bold=True, size=10)
body_para(doc,
    "Coastal fishermen operating in Cordova, Lapu-Lapu, Bantayan, Medellin, and Daanbantayan. "
    "These individuals represent the most underserved segment of Cebu's seafood economy—highest impact, "
    "lowest digital access. FreshCatch is designed with mobile-first UX to ensure usability even with "
    "basic smartphones."
)

body_para(doc, "Secondary Users — Restaurant, Resort, and Hotel Buyers", bold=True, size=10)
body_para(doc,
    "Cebu's food and hospitality sector—restaurants, resorts, and hotels—that procure seafood on a "
    "recurring basis. These buyers benefit from advance supply commitments, price transparency, and "
    "traceable sourcing that strengthens their Local & Sustainable marketing positions."
)

body_para(doc, "Tertiary Users — Cooperatives, LGUs, and Distributors", bold=True, size=10)
body_para(doc,
    "Fishing cooperatives that can aggregate supply from multiple fishermen. "
    "Local Government Units (LGUs) that want visibility into local seafood market data. "
    "Tourism establishments and food distributors aligned with sustainable supply chains."
)

# =============================================
# IV. PROPOSED SOLUTION AND EXECUTION
# =============================================
section_title(doc, "IV.", "Proposed Solution and Execution")

subsection_title(doc, "A.", "Proposed Solution")
body_para(doc,
    "FreshCatch is a digital B2B seafood coordination platform built on a simple but powerful mechanism: "
    "demand-first seafood coordination. Restaurants post what they need. Fishermen see what is committed "
    "before sailing. The platform aggregates demand, displays supply forecasts, and enables commitment "
    "transactions—without replacing the existing physical logistics layer."
)

features = [
    ("Catch Forecast Module", 
     "Fishermen post expected catch details: fish type, volume, landing date, suggested price, and location. "
     "This creates a live, browsable supply catalog for restaurant buyers."),
    ("Restaurant Demand Posting", 
     "Restaurants list their seafood requirements: species, quantity, delivery date, and preferred price range. "
     "Fishermen can see these publicly to guide what they target during a fishing trip."),
    ("Commitment System", 
     "A mutual confirmation layer where restaurants commit to purchase and fishermen confirm supply post-landing. "
     "This creates accountability and trust without complex financial escrow in the MVP."),
    ("Aggregated Demand Dashboard", 
     "Displays total committed volume per species (e.g., Tuna: 100kg expected, 70kg committed). "
     "Fishermen can gauge total market demand before deployment."),
    ("Price Transparency Layer", 
     "Both sides see current price ranges per species, reducing the information asymmetry "
     "that middlemen exploit."),
    ("Rating and Trust System", 
     "Restaurants rate suppliers. Fishermen build verifiable reputation. "
     "This enables sustainable B2B relationships and reduces repeat transaction friction."),
]

for name, desc in features:
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"{name}: ")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)

subsection_title(doc, "B.", "Execution")
body_para(doc,
    "The FreshCatch MVP will be built using a modern, scalable stack: Next.js (frontend), "
    "Supabase (backend-as-a-service with PostgreSQL), and deployed on Vercel for zero-downtime hosting. "
    "The MVP will be developed iteratively with three milestones:"
)

milestones = [
    "Week 1: Authentication system (fisherman / restaurant roles), catch posting, demand posting.",
    "Week 2: Commitment system, aggregated demand dashboard, price transparency layer.",
    "Week 3: Rating system, POC demo polish, UI refinement for pitch demonstration.",
]
for m in milestones:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(m)
    r.font.size = Pt(10)

body_para(doc,
    "The Proof of Concept (POC) prototype for the first pitch will demonstrate: user login with role selection, "
    "seafood demand posting, expected catch posting, and the aggregated demand dashboard with commitment simulation. "
    "Security is built in from the start via Supabase Row Level Security (RLS), ensuring fishermen and restaurants "
    "only see data relevant to their role. Optional AI features—demand prediction and price recommendations—"
    "can be layered in using OpenAI API in later sprints."
)

# =============================================
# REFERENCES
# =============================================
section_title(doc, "", "References")
refs = [
    "[1] Philippine Statistics Authority. (2023). Fisheries Situation Report. Manila, Philippines.",
    "[2] Department of Agriculture – Bureau of Fisheries and Aquatic Resources (BFAR). Cebu Fisheries Data, 2022.",
    "[3] Supabase Inc. (2024). Supabase Documentation. https://supabase.com/docs",
    "[4] Vercel Inc. (2024). Next.js Documentation. https://nextjs.org/docs",
    "[5] OpenAI. (2024). API Reference. https://platform.openai.com/docs",
]
for ref in refs:
    p = doc.add_paragraph()
    r = p.add_run(ref)
    r.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# =============================================
# V. FIVE YEAR ROAD-MAP
# =============================================
section_title(doc, "V.", "Five Year Road-Map")

roadmap_table = doc.add_table(rows=6, cols=3)
roadmap_table.style = 'Table Grid'

headers = ["Year", "Milestone", "Focus Area"]
header_row = roadmap_table.rows[0]
for i, h in enumerate(headers):
    cell = header_row.cells[i]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10)

roadmap_data = [
    ("1", 
     "Launch FreshCatch MVP in Cebu with 100+ active fishermen and 50+ restaurant partners. Secure seed funding or LGU partnership.", 
     "Product-market fit, user acquisition, Cebu coastal communities."),
    ("2", 
     "Introduce AI-powered demand prediction and pricing recommendations. Expand to Bohol and Negros Occidental.", 
     "Feature expansion and regional market penetration."),
    ("3", 
     "Integrate logistics partners for last-mile seafood delivery. Launch FreshCatch Cooperative Portal for aggregated supply.", 
     "Scalability, logistics integration, cooperative partnerships."),
    ("4", 
     "Attain revenue sustainability via transaction commissions and premium subscriptions. Begin development of FreshProduce module.", 
     "Financial viability, B2B market expansion, product diversification."),
    ("5", 
     "Become the leading digital food supply chain platform for Cebu. Explore Series A funding and IPO readiness.", 
     "Market dominance, full food supply chain, national expansion."),
]

for i, (year, milestone, focus) in enumerate(roadmap_data):
    row = roadmap_table.rows[i + 1]
    row.cells[0].text = year
    row.cells[1].text = milestone
    row.cells[2].text = focus
    for cell in row.cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.size = Pt(10)

# =============================================
# SAVE
# =============================================
output_path = r"C:\Users\subar\OneDrive\Desktop\Hackathon Pitch\CFM_VPD.docx"
doc.save(output_path)
print(f"✅ VPD saved to: {output_path}")
